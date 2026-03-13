#!/usr/bin/env python3
"""
Genera el informe de protocolo de uso de la aplicación MedLoan
en formato Word (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                            font_size=None, color=None, alignment=None, space_after=None, space_before=None):
    """Añade un párrafo con formato personalizado."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def create_styled_table(doc, headers, rows, col_widths=None):
    """Crea una tabla con estilo profesional."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "0D9488")  # Teal

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0FDFA")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def generate_report():
    doc = Document()

    # ==================== ESTILOS ====================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(13, 148, 136)  # Teal

    # ==================== PORTADA ====================
    for _ in range(6):
        doc.add_paragraph('')

    add_formatted_paragraph(doc, 'HOSPITAL UNIVERSITARIO DE FUENLABRADA',
                          font_size=14, bold=True, color=(100, 100, 100),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, 'Servicio de Farmacia',
                          font_size=12, italic=True, color=(120, 120, 120),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(13, 148, 136)
    run.font.size = Pt(14)

    add_formatted_paragraph(doc, 'PROTOCOLO DE USO',
                          font_size=28, bold=True, color=(13, 148, 136),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, 'Sistema de Gestión de Préstamos',
                          font_size=20, bold=True, color=(50, 50, 50),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'de Medicamentos Interhospitalarios',
                          font_size=20, bold=True, color=(50, 50, 50),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, 'MedLoan',
                          font_size=16, italic=True, color=(13, 148, 136),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(13, 148, 136)
    run.font.size = Pt(14)

    add_formatted_paragraph(doc, 'Versión 1.0 — Febrero 2026',
                          font_size=11, color=(130, 130, 130),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'Documento de uso interno',
                          font_size=10, italic=True, color=(150, 150, 150),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ==================== ÍNDICE ====================
    doc.add_heading('ÍNDICE DE CONTENIDOS', level=1)

    toc_items = [
        ('1.', 'Introducción y Objetivo del Sistema'),
        ('2.', 'Acceso al Sistema'),
        ('   2.1.', 'Inicio de Sesión'),
        ('   2.2.', 'Cierre de Sesión'),
        ('3.', 'Panel de Control (Dashboard)'),
        ('   3.1.', 'Indicadores Clave (KPIs)'),
        ('   3.2.', 'Actividad Reciente'),
        ('   3.3.', 'Accesos Rápidos'),
        ('4.', 'Gestión de Préstamos'),
        ('   4.1.', 'Tipos de Préstamos'),
        ('   4.2.', 'Crear un Nuevo Préstamo'),
        ('   4.3.', 'Listado de Préstamos'),
        ('   4.4.', 'Filtros y Búsqueda'),
        ('   4.5.', 'Cambiar Estado de un Préstamo'),
        ('   4.6.', 'Operaciones Masivas'),
        ('   4.7.', 'Detalle de un Préstamo'),
        ('5.', 'Gestión de Hospitales'),
        ('   5.1.', 'Crear un Hospital'),
        ('   5.2.', 'Editar y Desactivar Hospitales'),
        ('6.', 'Gestión de Medicamentos'),
        ('   6.1.', 'Catálogo de Medicamentos'),
        ('   6.2.', 'Búsqueda CIMA'),
        ('   6.3.', 'Crear Medicamento Manualmente'),
        ('7.', 'Generación de Documentos PDF'),
        ('   7.1.', 'PDF de Préstamo Individual'),
        ('   7.2.', 'PDF de Listado de Pendientes'),
        ('8.', 'Notificaciones por Correo Electrónico'),
        ('   8.1.', 'Envío Automático'),
        ('   8.2.', 'Contenido del Correo'),
        ('9.', 'Estadísticas y Analítica'),
        ('   9.1.', 'Indicadores y Gráficas'),
        ('   9.2.', 'Exportación de Informes'),
        ('10.', 'Configuración del Sistema'),
        ('   10.1.', 'Configuración SMTP'),
        ('   10.2.', 'Email en Copia (CC)'),
        ('11.', 'Flujo de Trabajo Completo'),
        ('   11.1.', 'Flujo: Solicitamos Medicamento'),
        ('   11.2.', 'Flujo: Nos Solicitan Medicamento'),
        ('12.', 'Roles y Responsabilidades'),
        ('13.', 'Glosario de Términos'),
        ('14.', 'Preguntas Frecuentes (FAQ)'),
        ('15.', 'Soporte y Contacto'),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run_num = p.add_run(num + ' ')
        run_num.bold = True
        run_num.font.size = Pt(10)
        if num.startswith('   '):
            run_num.font.color.rgb = RGBColor(80, 80, 80)
        else:
            run_num.font.color.rgb = RGBColor(13, 148, 136)
        run_title = p.add_run(title)
        run_title.font.size = Pt(10)

    doc.add_page_break()

    # ==================== 1. INTRODUCCIÓN ====================
    doc.add_heading('1. Introducción y Objetivo del Sistema', level=1)

    doc.add_paragraph(
        'MedLoan es el sistema de gestión de préstamos de medicamentos interhospitalarios '
        'del Hospital Universitario de Fuenlabrada (HUF). Esta aplicación web ha sido '
        'diseñada específicamente para el Servicio de Farmacia con el objetivo de digitalizar, '
        'centralizar y optimizar todo el proceso de préstamo y devolución de medicamentos '
        'entre hospitales del sistema sanitario.'
    )

    doc.add_heading('Objetivos principales:', level=3)

    objectives = [
        'Registrar y documentar todos los préstamos de medicamentos realizados entre hospitales, '
        'tanto los que solicita el HUF como los que solicitan otros hospitales al HUF.',
        'Generar documentación oficial en formato PDF con la imagen corporativa del hospital, '
        'incluyendo todos los datos necesarios del préstamo (medicamentos, cantidades, hospital, etc.).',
        'Automatizar el envío de notificaciones por correo electrónico al hospital proveedor '
        'cuando se solicita un préstamo, adjuntando el PDF del préstamo.',
        'Controlar el estado de cada préstamo mediante dos indicadores independientes: '
        'si ha sido gestionado en Farmatools (el sistema de gestión farmacéutico del hospital) '
        'y si el medicamento ha sido devuelto.',
        'Proporcionar estadísticas y análisis detallados sobre la actividad de préstamos, '
        'permitiendo identificar tendencias, hospitales más frecuentes y medicamentos más solicitados.',
        'Mantener un catálogo actualizado de hospitales y medicamentos, con integración con la base '
        'de datos CIMA de la Agencia Española de Medicamentos y Productos Sanitarios (AEMPS).',
    ]

    for obj in objectives:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        'El sistema está accesible a través de un navegador web y está protegido mediante '
        'contraseña. No requiere instalación de software adicional en los equipos de los usuarios.'
    )

    doc.add_page_break()

    # ==================== 2. ACCESO AL SISTEMA ====================
    doc.add_heading('2. Acceso al Sistema', level=1)

    doc.add_heading('2.1. Inicio de Sesión', level=2)

    doc.add_paragraph(
        'Para acceder al sistema, abra un navegador web (se recomienda Google Chrome o '
        'Microsoft Edge) y acceda a la URL del sistema proporcionada por el administrador.'
    )

    doc.add_heading('Pasos para iniciar sesión:', level=3)

    login_steps = [
        'Abra el navegador web e introduzca la dirección URL del sistema.',
        'Se mostrará la pantalla de inicio de sesión con el logotipo del hospital.',
        'Introduzca la contraseña proporcionada por el administrador del sistema en el campo "Contraseña".',
        'Pulse el botón "Entrar" o presione la tecla Enter.',
        'Si la contraseña es correcta, será redirigido automáticamente al Panel de Control (Dashboard).',
        'Si la contraseña es incorrecta, se mostrará un mensaje de error. Verifique la contraseña e intente de nuevo.',
    ]

    for i, step in enumerate(login_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Paso {i}: ')
        run.bold = True
        p.add_run(step)
        p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('Nota importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 80, 0)
    p.add_run(
        'La sesión permanece activa durante 30 días. Pasado este tiempo, '
        'deberá volver a introducir la contraseña. Todos los usuarios comparten la misma '
        'contraseña de acceso, ya que el sistema está diseñado para uso exclusivo del '
        'Servicio de Farmacia.'
    )

    doc.add_heading('2.2. Cierre de Sesión', level=2)

    doc.add_paragraph(
        'Para cerrar la sesión, haga clic en el botón "Cerrar sesión" situado en la barra '
        'lateral de navegación (sidebar), en la parte inferior. Al cerrar sesión, será '
        'redirigido a la pantalla de inicio de sesión y deberá introducir la contraseña '
        'nuevamente para volver a acceder.'
    )

    doc.add_page_break()

    # ==================== 3. PANEL DE CONTROL ====================
    doc.add_heading('3. Panel de Control (Dashboard)', level=1)

    doc.add_paragraph(
        'El Panel de Control es la pantalla principal del sistema. Proporciona una visión '
        'general del estado actual de los préstamos y permite acceder rápidamente a las '
        'funciones más utilizadas. Se accede automáticamente tras iniciar sesión o pulsando '
        '"Dashboard" en el menú lateral.'
    )

    doc.add_heading('3.1. Indicadores Clave (KPIs)', level=2)

    doc.add_paragraph(
        'En la parte superior del Dashboard se muestran tarjetas con los indicadores clave de rendimiento:'
    )

    kpi_data = [
        ['Préstamos activos', 'Número total de préstamos que están actualmente en curso (no devueltos).'],
        ['Pendientes Farmatools', 'Préstamos que aún no han sido registrados/gestionados en el sistema Farmatools.'],
        ['Pendientes de devolver', 'Préstamos de tipo "Solicitamos" cuyo medicamento aún no ha sido devuelto al hospital proveedor.'],
        ['Pendientes de que nos devuelvan', 'Préstamos de tipo "Nos solicitan" cuyo medicamento aún no nos ha sido devuelto por el hospital solicitante.'],
    ]

    create_styled_table(doc, ['Indicador', 'Descripción'], kpi_data, [5, 12])

    doc.add_heading('3.2. Actividad Reciente', level=2)

    doc.add_paragraph(
        'Debajo de los KPIs se muestra una lista con los 8 préstamos más recientes, '
        'incluyendo el número de referencia, el tipo de préstamo (Solicitamos / Nos solicitan), '
        'el hospital implicado, la fecha y los estados actuales. Esto permite al farmacéutico '
        'tener una visión rápida de la actividad más reciente sin necesidad de navegar al listado completo.'
    )

    doc.add_heading('3.3. Accesos Rápidos', level=2)

    doc.add_paragraph(
        'El Dashboard incluye botones de acción rápida que permiten:'
    )

    quick_actions = [
        'Crear un nuevo préstamo directamente.',
        'Ver el listado completo de préstamos.',
        'Acceder a las estadísticas.',
        'Visualizar la tendencia mensual de préstamos mediante un gráfico sparkline.',
    ]

    for action in quick_actions:
        doc.add_paragraph(action, style='List Bullet')

    doc.add_page_break()

    # ==================== 4. GESTIÓN DE PRÉSTAMOS ====================
    doc.add_heading('4. Gestión de Préstamos', level=1)

    doc.add_paragraph(
        'La gestión de préstamos es la funcionalidad central del sistema. Permite crear, '
        'consultar, filtrar y gestionar todos los préstamos de medicamentos entre hospitales.'
    )

    doc.add_heading('4.1. Tipos de Préstamos', level=2)

    doc.add_paragraph(
        'El sistema distingue dos tipos de préstamos, según la dirección del flujo de medicamentos:'
    )

    # Tabla de tipos
    type_data = [
        ['SOLICITAMOS\n(Nosotros solicitamos)',
         'El HUF solicita un medicamento a otro hospital. '
         'El HUF es el receptor del medicamento y deberá devolverlo posteriormente.\n\n'
         'Al crear este tipo de préstamo, el sistema envía automáticamente un correo '
         'electrónico al hospital proveedor con el PDF del préstamo adjunto.',
         'Morado / Púrpura'],
        ['NOS SOLICITAN\n(Otro hospital nos solicita)',
         'Otro hospital solicita un medicamento al HUF. '
         'El HUF es el proveedor del medicamento y espera que se lo devuelvan.\n\n'
         'Este tipo de préstamo NO envía correo electrónico automáticamente, ya que '
         'se asume que el hospital solicitante ya ha contactado previamente.',
         'Verde azulado (Teal)'],
    ]

    create_styled_table(doc, ['Tipo', 'Descripción', 'Color identificativo'], type_data, [4, 10, 3])

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 80, 0)
    p.add_run(
        'Es fundamental seleccionar correctamente el tipo de préstamo, ya que determina '
        'el flujo de trabajo posterior (envío de email, etiquetas de estado, etc.).'
    )

    doc.add_heading('4.2. Crear un Nuevo Préstamo', level=2)

    doc.add_paragraph(
        'Para crear un nuevo préstamo, acceda a Préstamos > Nuevo Préstamo desde el menú lateral, '
        'o utilice el botón "+ Nuevo Préstamo" disponible en el listado de préstamos o en el Dashboard.'
    )

    doc.add_heading('Campos del formulario:', level=3)

    form_fields = [
        ['Tipo de préstamo *', 'Seleccione "Solicitamos" o "Nos solicitan". Este campo es obligatorio y determina '
         'el comportamiento del sistema.'],
        ['Hospital *', 'Seleccione el hospital implicado en el préstamo del listado desplegable. Si el hospital '
         'no existe, deberá crearlo primero en la sección de Hospitales.'],
        ['Medicamentos *', 'Busque y seleccione uno o varios medicamentos. Puede buscar por nombre, código nacional (CN) '
         'o principio activo. Si el medicamento no existe en el catálogo, puede crearlo directamente desde el '
         'formulario o buscarlo en CIMA.'],
        ['Unidades *', 'Para cada medicamento seleccionado, indique el número de unidades prestadas. Debe ser un '
         'número entero mayor que 0.'],
        ['Nombre del farmacéutico', 'Nombre del farmacéutico que gestiona el préstamo. Campo opcional pero recomendado '
         'para trazabilidad.'],
        ['Email de envío', 'Dirección de correo electrónico a la que se enviará la notificación con el PDF adjunto. '
         'Solo aplica para préstamos de tipo "Solicitamos". Se pre-rellena con el email del hospital si está disponible.'],
        ['Notas / Observaciones', 'Campo de texto libre para añadir cualquier información adicional relevante '
         'sobre el préstamo (motivo, urgencia, etc.). Campo opcional.'],
    ]

    create_styled_table(doc, ['Campo', 'Descripción'], form_fields, [4.5, 12.5])

    doc.add_paragraph('')
    doc.add_heading('Proceso de creación paso a paso:', level=3)

    creation_steps = [
        ('Seleccionar tipo de préstamo',
         'Haga clic en "Solicitamos" o "Nos solicitan" según corresponda.'),
        ('Seleccionar hospital',
         'Utilice el desplegable para buscar y seleccionar el hospital. Puede escribir para filtrar.'),
        ('Añadir medicamentos',
         'En el campo de búsqueda de medicamentos, escriba el nombre del medicamento, el código nacional '
         '(CN) o el principio activo. Seleccione el medicamento de los resultados. Si no existe, pulse '
         '"Crear nuevo" para añadirlo al catálogo, o use "Buscar en CIMA" para importarlo desde la base '
         'de datos de la AEMPS.'),
        ('Indicar unidades',
         'Para cada medicamento añadido, introduzca el número de unidades en el campo correspondiente.'),
        ('Completar datos opcionales',
         'Rellene el nombre del farmacéutico, el email de envío (si es tipo "Solicitamos") y las notas '
         'que considere necesarias.'),
        ('Guardar el préstamo',
         'Pulse el botón "Guardar préstamo". El sistema:\n'
         '   a) Generará automáticamente un número de referencia único (formato PREST-AAAA-NNNNN).\n'
         '   b) Creará el registro en la base de datos.\n'
         '   c) Generará el PDF del préstamo.\n'
         '   d) Si es tipo "Solicitamos", enviará el email al hospital con el PDF adjunto.\n'
         '   e) Redirigirá al listado de préstamos con un mensaje de confirmación.'),
    ]

    for i, (title, desc) in enumerate(creation_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Paso {i} — {title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    doc.add_heading('4.3. Listado de Préstamos', level=2)

    doc.add_paragraph(
        'El listado de préstamos muestra todos los préstamos registrados en el sistema en formato '
        'de tabla. Se accede desde el menú lateral pulsando "Préstamos".'
    )

    doc.add_heading('Columnas de la tabla:', level=3)

    columns_data = [
        ['Selección (checkbox)', 'Casilla de selección para operaciones masivas.'],
        ['Referencia', 'Número de referencia único del préstamo (ej: PREST-2026-00042).'],
        ['Tipo', 'Etiqueta visual que indica si es "Solicitamos" (morado) o "Nos solicitan" (teal).'],
        ['Hospital', 'Nombre del hospital implicado en el préstamo.'],
        ['Medicamentos', 'Lista de medicamentos incluidos en el préstamo con sus unidades.'],
        ['Fecha', 'Fecha y hora de creación del préstamo.'],
        ['Farmatools', 'Toggle/interruptor que indica si el préstamo ha sido gestionado en Farmatools. '
         'Se puede cambiar directamente desde la tabla.'],
        ['Devuelto', 'Toggle/interruptor que indica si el medicamento ha sido devuelto. '
         'Se puede cambiar directamente desde la tabla.'],
        ['Acciones', 'Botones para ver detalle, descargar PDF o eliminar el préstamo.'],
    ]

    create_styled_table(doc, ['Columna', 'Descripción'], columns_data, [4, 13])

    doc.add_paragraph('')
    doc.add_heading('4.4. Filtros y Búsqueda', level=2)

    doc.add_paragraph(
        'El listado de préstamos dispone de un sistema completo de filtros para localizar rápidamente '
        'los préstamos que necesite:'
    )

    filters_data = [
        ['Búsqueda de texto', 'Busca simultáneamente en: número de referencia, nombre del hospital y '
         'nombre de los medicamentos. Escriba cualquier término para filtrar.'],
        ['Filtro por tipo', 'Permite mostrar solo préstamos de tipo "Solicitamos", solo "Nos solicitan", '
         'o todos.'],
        ['Filtro Farmatools', 'Filtra por estado de gestión en Farmatools: Todos, Gestionado, o Pendiente.'],
        ['Filtro Devolución', 'Filtra por estado de devolución: Todos, Devuelto, o Pendiente.'],
        ['Filtro por hospital', 'Permite seleccionar un hospital específico para ver solo sus préstamos.'],
    ]

    create_styled_table(doc, ['Filtro', 'Descripción'], filters_data, [4.5, 12.5])

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Consejo: ')
    run.bold = True
    run.font.color.rgb = RGBColor(13, 148, 136)
    p.add_run(
        'Los filtros se pueden combinar entre sí. Por ejemplo, puede filtrar para ver solo '
        'los préstamos de tipo "Solicitamos" que estén pendientes de devolución para un hospital concreto.'
    )

    doc.add_heading('4.5. Cambiar Estado de un Préstamo', level=2)

    doc.add_paragraph(
        'Cada préstamo tiene dos estados independientes que se pueden modificar directamente '
        'desde el listado sin necesidad de acceder al detalle:'
    )

    doc.add_heading('Estado "Farmatools":', level=3)
    doc.add_paragraph(
        'Indica si el préstamo ha sido registrado/gestionado en el sistema Farmatools del hospital. '
        'Para cambiar el estado, haga clic en el toggle (interruptor) de la columna "Farmatools" en la '
        'fila del préstamo correspondiente. El cambio se aplica de forma instantánea (actualización optimista) '
        'y se confirma con una notificación visual.'
    )

    doc.add_heading('Estado "Devuelto":', level=3)
    doc.add_paragraph(
        'Indica si el medicamento prestado ha sido devuelto (o nos lo han devuelto, según el tipo). '
        'Para cambiar el estado, haga clic en el toggle de la columna "Devuelto". El cambio se aplica '
        'inmediatamente. Cuando un préstamo está marcado como devuelto, se considera completado.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Nota: ')
    run.bold = True
    p.add_run(
        'Ambos estados son independientes. Un préstamo puede estar gestionado en Farmatools pero '
        'no devuelto, o viceversa. El flujo habitual es: primero se gestiona en Farmatools y luego '
        'se marca como devuelto cuando se produce la devolución física.'
    )

    doc.add_heading('4.6. Operaciones Masivas', level=2)

    doc.add_paragraph(
        'El sistema permite realizar operaciones sobre múltiples préstamos a la vez:'
    )

    bulk_ops = [
        ('Selección múltiple',
         'Marque las casillas de verificación (checkboxes) de los préstamos que desee. '
         'También puede usar la casilla de la cabecera para seleccionar/deseleccionar todos los '
         'préstamos visibles.'),
        ('Marcar como devueltos',
         'Seleccione los préstamos deseados y pulse "Marcar como devueltos". Todos los '
         'préstamos seleccionados cambiarán su estado de devolución a "Devuelto".'),
        ('Eliminar préstamos',
         'Seleccione los préstamos deseados y pulse "Eliminar seleccionados". Se mostrará un '
         'diálogo de confirmación antes de proceder. Esta acción es irreversible y eliminará '
         'permanentemente los préstamos seleccionados junto con todos sus elementos (medicamentos).'),
    ]

    for title, desc in bulk_ops:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('4.7. Detalle de un Préstamo', level=2)

    doc.add_paragraph(
        'Para ver todos los detalles de un préstamo, haga clic en el número de referencia '
        'o en el botón de "Ver detalle" en la fila correspondiente. La página de detalle muestra:'
    )

    detail_items = [
        'Número de referencia y fecha de creación.',
        'Tipo de préstamo con etiqueta visual.',
        'Hospital implicado con sus datos de contacto.',
        'Lista completa de medicamentos con código nacional, presentación, principio activo y unidades.',
        'Nombre del farmacéutico responsable.',
        'Notas/observaciones del préstamo.',
        'Estado actual (Farmatools y devolución).',
        'Email al que se envió la notificación (si aplica).',
        'Opciones para descargar el PDF, editar notas o cambiar estados.',
    ]

    for item in detail_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==================== 5. GESTIÓN DE HOSPITALES ====================
    doc.add_heading('5. Gestión de Hospitales', level=1)

    doc.add_paragraph(
        'La sección de Hospitales permite mantener el catálogo de hospitales con los que '
        'el HUF realiza préstamos de medicamentos. Se accede desde el menú lateral pulsando "Hospitales".'
    )

    doc.add_heading('5.1. Crear un Hospital', level=2)

    doc.add_paragraph('Para dar de alta un nuevo hospital, pulse el botón "+ Nuevo Hospital" y complete los siguientes campos:')

    hospital_fields = [
        ['Nombre *', 'Nombre completo del hospital (ej: "Hospital Universitario 12 de Octubre"). Campo obligatorio.'],
        ['Email', 'Dirección de correo electrónico del Servicio de Farmacia del hospital. Este email se utilizará '
         'como valor predeterminado al crear préstamos de tipo "Solicitamos".'],
        ['Dirección', 'Dirección física del hospital. Campo opcional.'],
        ['Teléfono', 'Número de teléfono de contacto. Campo opcional.'],
    ]

    create_styled_table(doc, ['Campo', 'Descripción'], hospital_fields, [4, 13])

    doc.add_paragraph('')
    doc.add_heading('5.2. Editar y Desactivar Hospitales', level=2)

    doc.add_paragraph(
        'Desde el listado de hospitales puede:'
    )

    hospital_ops = [
        ('Editar', 'Modifique cualquier dato del hospital pulsando el botón de edición. Los cambios se reflejarán '
         'en futuros préstamos, pero no modificarán préstamos ya existentes.'),
        ('Desactivar', 'En lugar de eliminar un hospital, el sistema lo desactiva (borrado lógico). '
         'Esto preserva la integridad de los préstamos históricos asociados. Un hospital desactivado no '
         'aparecerá en los desplegables de selección al crear nuevos préstamos, pero sus préstamos '
         'existentes seguirán siendo visibles.'),
        ('Reactivar', 'Un hospital desactivado puede reactivarse en cualquier momento, restaurando su '
         'disponibilidad para nuevos préstamos.'),
    ]

    for title, desc in hospital_ops:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ==================== 6. GESTIÓN DE MEDICAMENTOS ====================
    doc.add_heading('6. Gestión de Medicamentos', level=1)

    doc.add_paragraph(
        'La sección de Medicamentos permite gestionar el catálogo de medicamentos disponibles '
        'para préstamos. Se accede desde el menú lateral pulsando "Medicamentos".'
    )

    doc.add_heading('6.1. Catálogo de Medicamentos', level=2)

    doc.add_paragraph(
        'El catálogo muestra todos los medicamentos registrados en el sistema. Cada medicamento '
        'contiene la siguiente información:'
    )

    med_fields = [
        ['Nombre *', 'Nombre comercial o descripción del medicamento. Campo obligatorio.'],
        ['Código Nacional (CN)', 'Código Nacional asignado por la AEMPS. Sirve como identificador único '
         'del medicamento en el sistema sanitario español.'],
        ['Presentación', 'Forma farmacéutica del medicamento (comprimidos, cápsulas, viales, ampollas, etc.).'],
        ['Principio activo', 'Nombre del compuesto activo del medicamento.'],
    ]

    create_styled_table(doc, ['Campo', 'Descripción'], med_fields, [4.5, 12.5])

    doc.add_paragraph('')
    doc.add_heading('6.2. Búsqueda CIMA', level=2)

    doc.add_paragraph(
        'El sistema está integrado con CIMA (Centro de Información de Medicamentos Autorizados), '
        'la base de datos oficial de la Agencia Española de Medicamentos y Productos Sanitarios (AEMPS). '
        'Esta integración permite buscar e importar medicamentos directamente al catálogo.'
    )

    doc.add_heading('¿Cómo funciona la búsqueda CIMA?', level=3)

    cima_steps = [
        'En el formulario de creación de préstamos, al buscar un medicamento que no exista en el catálogo local, '
        'aparecerá la opción "Buscar en CIMA".',
        'Se puede buscar por tres criterios: código nacional (CN), nombre del medicamento o principio activo.',
        'La búsqueda por principio activo tiene prioridad y los resultados se muestran primero.',
        'Se muestran hasta 15 resultados por búsqueda.',
        'Al seleccionar un medicamento de los resultados CIMA, este se importa automáticamente al catálogo local '
        'con todos sus datos (nombre, CN, presentación, principio activo).',
        'Si el medicamento ya existe en el catálogo (mismo CN o nombre), se reutiliza el existente en lugar de crear un duplicado.',
    ]

    for i, step in enumerate(cima_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. ')
        run.bold = True
        p.add_run(step)
        p.paragraph_format.space_after = Pt(3)

    doc.add_heading('6.3. Crear Medicamento Manualmente', level=3)

    doc.add_paragraph(
        'Si un medicamento no se encuentra en CIMA o desea añadirlo manualmente, puede hacerlo '
        'de dos formas:'
    )

    manual_med = [
        ('Desde la sección Medicamentos',
         'Pulse "+ Nuevo Medicamento" y complete los campos del formulario.'),
        ('Desde el formulario de préstamo',
         'Al buscar un medicamento y no encontrar resultados, pulse "Crear nuevo". '
         'El medicamento se creará automáticamente y se añadirá al préstamo.'),
    ]

    for title, desc in manual_med:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        'Al igual que con los hospitales, los medicamentos se desactivan en lugar de eliminarse '
        '(borrado lógico), preservando la integridad de los préstamos históricos.'
    )

    doc.add_page_break()

    # ==================== 7. GENERACIÓN DE DOCUMENTOS PDF ====================
    doc.add_heading('7. Generación de Documentos PDF', level=1)

    doc.add_paragraph(
        'El sistema genera documentos PDF profesionales con la imagen corporativa del '
        'Hospital Universitario de Fuenlabrada. Existen dos tipos de documentos PDF:'
    )

    doc.add_heading('7.1. PDF de Préstamo Individual', level=2)

    doc.add_paragraph(
        'Cada préstamo tiene asociado un documento PDF que se genera automáticamente al crear '
        'el préstamo. Este PDF puede descargarse en cualquier momento desde:'
    )

    pdf_access = [
        'El listado de préstamos, pulsando el icono de descarga en la columna de acciones.',
        'La página de detalle del préstamo, pulsando el botón "Descargar PDF".',
    ]

    for item in pdf_access:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Contenido del PDF de préstamo:', level=3)

    pdf_content = [
        'Logotipo del Hospital Universitario de Fuenlabrada en la cabecera.',
        'Título del documento con el tipo de préstamo identificado por color (morado para "Solicitamos", teal para "Nos solicitan").',
        'Número de referencia único del préstamo.',
        'Fecha y hora de creación.',
        'Datos del hospital implicado.',
        'Nombre del farmacéutico responsable.',
        'Tabla de medicamentos con: nombre, código nacional, presentación, principio activo y unidades. '
        'Si solo hay un medicamento, se muestra en formato de campos detallados. Si hay varios, se muestra '
        'en formato tabla con una fila de totales.',
        'Sección de notas/observaciones (si las hubiera).',
        'Líneas de firma para ambos servicios de farmacia (el del HUF y el del hospital colaborador).',
        'Pie de página con la dirección y datos de contacto del HUF.',
    ]

    for item in pdf_content:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2. PDF de Listado de Pendientes', level=2)

    doc.add_paragraph(
        'El sistema puede generar un informe PDF con el listado de todos los préstamos pendientes '
        'de devolución. Este documento es útil para revisar periódicamente el estado de las devoluciones.'
    )

    doc.add_heading('Tipos de listado de pendientes:', level=3)

    pending_types = [
        ('Pendientes de devolver',
         'Muestra los préstamos de tipo "Solicitamos" que aún no han sido devueltos al hospital '
         'proveedor. Es decir, medicamentos que el HUF tiene prestados y debe devolver.'),
        ('Pendientes de que nos devuelvan',
         'Muestra los préstamos de tipo "Nos solicitan" que aún no nos han devuelto. Es decir, '
         'medicamentos que el HUF ha prestado a otros hospitales y espera que le devuelvan.'),
    ]

    for title, desc in pending_types:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'Estos PDF se generan en formato apaisado (landscape) y contienen una tabla con las columnas: '
        'Referencia, Fecha, Medicamento, Hospital y Unidades. Al final de la tabla se muestra una fila '
        'con los totales. El documento soporta múltiples páginas con cabeceras y pies de página automáticos.'
    )

    doc.add_page_break()

    # ==================== 8. NOTIFICACIONES POR EMAIL ====================
    doc.add_heading('8. Notificaciones por Correo Electrónico', level=1)

    doc.add_heading('8.1. Envío Automático', level=2)

    doc.add_paragraph(
        'Cuando se crea un préstamo de tipo "Solicitamos", el sistema envía automáticamente un correo '
        'electrónico al hospital proveedor con el PDF del préstamo adjunto. Este envío se realiza '
        'de forma transparente durante el proceso de creación del préstamo.'
    )

    doc.add_heading('Condiciones del envío:', level=3)

    email_conditions = [
        'Solo se envía para préstamos de tipo "Solicitamos" (NO para "Nos solicitan").',
        'Es necesario indicar una dirección de email de destino en el formulario de creación.',
        'El sistema SMTP debe estar correctamente configurado (ver sección 10).',
        'Si el envío falla, el préstamo se crea igualmente pero se notifica al usuario del error.',
        'Si se ha configurado un email en copia (CC), también se enviará una copia del correo a esa dirección.',
    ]

    for item in email_conditions:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2. Contenido del Correo', level=3)

    doc.add_paragraph('El correo electrónico incluye:')

    email_content = [
        ('Saludo personalizado', 'Un saludo que varía según la hora del día en España:\n'
         '   — "Buenos días" (de 6:00 a 13:59)\n'
         '   — "Buenas tardes" (de 14:00 en adelante)'),
        ('Cuerpo del mensaje',
         'Un texto informativo que indica que se adjunta la documentación relativa al préstamo '
         'de medicamentos solicitado, incluyendo el número de referencia.'),
        ('PDF adjunto', 'El documento PDF completo del préstamo como archivo adjunto.'),
        ('Remitente', 'El nombre y email configurados en la sección de Configuración SMTP.'),
        ('Copia (CC)', 'Si se ha configurado un email en CC en los ajustes SMTP, se enviará '
         'una copia del correo a esa dirección (útil para que la secretaria o responsable '
         'mantenga un registro).'),
    ]

    for title, desc in email_content:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ==================== 9. ESTADÍSTICAS Y ANALÍTICA ====================
    doc.add_heading('9. Estadísticas y Analítica', level=1)

    doc.add_paragraph(
        'La sección de Estadísticas proporciona un análisis detallado de la actividad de préstamos '
        'del servicio. Se accede desde el menú lateral pulsando "Estadísticas".'
    )

    doc.add_heading('9.1. Indicadores y Gráficas', level=2)

    doc.add_heading('Filtros de fecha:', level=3)
    doc.add_paragraph(
        'Todas las estadísticas se pueden filtrar por periodo temporal. Los filtros predefinidos son:'
    )

    date_filters = [
        'Últimos 7 días',
        'Últimos 30 días',
        'Últimos 3 meses',
        'Último año',
        'Rango personalizado (seleccione fecha de inicio y fin)',
    ]

    for f in date_filters:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph(
        'Además, se puede configurar un umbral de días para considerar un préstamo como "atrasado" '
        '(por defecto, 30 días sin devolución).'
    )

    doc.add_heading('KPIs principales:', level=3)

    stats_kpis = [
        ['Total de préstamos', 'Número total de préstamos en el periodo seleccionado.'],
        ['Total de unidades', 'Suma de todas las unidades de medicamentos prestadas.'],
        ['Media de unidades por préstamo', 'Promedio de unidades por cada préstamo individual.'],
        ['Tasa de devolución', 'Porcentaje de préstamos que ya han sido devueltos sobre el total.'],
    ]

    create_styled_table(doc, ['KPI', 'Descripción'], stats_kpis, [5.5, 11.5])

    doc.add_paragraph('')
    doc.add_heading('Gráficos disponibles:', level=3)

    charts = [
        ('Volumen de préstamos en el tiempo',
         'Gráfico de líneas/área que muestra la evolución del número de préstamos a lo largo del tiempo, '
         'diferenciando entre "Solicitamos" y "Nos solicitan". Permite identificar tendencias y picos '
         'de actividad.'),
        ('Distribución por tipo',
         'Gráfico circular (donut) que muestra la proporción entre préstamos solicitados y recibidos.'),
        ('Top hospitales',
         'Gráfico de barras horizontales con los hospitales que más préstamos han tenido, ordenados '
         'de mayor a menor.'),
        ('Top medicamentos',
         'Gráfico de barras horizontales con los medicamentos más frecuentemente prestados.'),
    ]

    for title, desc in charts:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('Tablas de análisis:', level=3)

    analysis_tables = [
        ('Ranking de hospitales',
         'Tabla ordenada con: nombre del hospital, número de préstamos, total de unidades y '
         'préstamos pendientes de devolución.'),
        ('Top medicamentos',
         'Tabla con los medicamentos más prestados, incluyendo número de veces y unidades totales.'),
        ('Préstamos atrasados',
         'Tabla detallada de préstamos que han superado el umbral de días configurado sin ser devueltos. '
         'Incluye referencia, hospital, medicamento, fecha y días transcurridos.'),
    ]

    for title, desc in analysis_tables:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('9.2. Exportación de Informes', level=2)

    doc.add_paragraph('Las estadísticas se pueden exportar en dos formatos:')

    export_formats = [
        ('Informe PDF',
         'Genera un documento PDF profesional con todos los KPIs, gráficos y tablas del periodo '
         'seleccionado. Ideal para compartir con la dirección o adjuntar a informes de actividad.'),
        ('Libro Excel (XLSX)',
         'Genera un archivo Excel con múltiples hojas:\n'
         '   — Resumen: KPIs principales.\n'
         '   — Volumen temporal: Datos del gráfico de evolución.\n'
         '   — Hospitales: Ranking completo con datos de préstamos.\n'
         '   — Medicamentos: Listado de medicamentos más prestados.\n'
         '   — Distribución por tipo: Datos del gráfico circular.\n'
         '   — Préstamos atrasados: Detalle de préstamos fuera de plazo (si los hubiera).'),
    ]

    for title, desc in export_formats:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ==================== 10. CONFIGURACIÓN ====================
    doc.add_heading('10. Configuración del Sistema', level=1)

    doc.add_paragraph(
        'La sección de Configuración permite ajustar los parámetros de envío de correo electrónico. '
        'Se accede desde el menú lateral pulsando "Configuración".'
    )

    doc.add_heading('10.1. Configuración SMTP', level=2)

    doc.add_paragraph(
        'Para que el sistema pueda enviar correos electrónicos con los PDFs de los préstamos, '
        'es necesario configurar los datos del servidor SMTP:'
    )

    smtp_fields = [
        ['Servidor SMTP (Host)', 'Dirección del servidor de correo (ej: smtp.gmail.com, smtp.office365.com).'],
        ['Puerto', 'Puerto del servidor SMTP. Valores habituales: 587 (TLS), 465 (SSL), 25 (sin cifrado).'],
        ['Conexión segura (TLS/SSL)', 'Active esta opción si el servidor requiere conexión cifrada (recomendado).'],
        ['Usuario', 'Nombre de usuario para la autenticación SMTP (normalmente la dirección de email).'],
        ['Contraseña', 'Contraseña o contraseña de aplicación para la autenticación SMTP.'],
        ['Nombre del remitente', 'Nombre que aparecerá como remitente del correo '
         '(ej: "Farmacia HU Fuenlabrada").'],
        ['Email del remitente', 'Dirección de correo que aparecerá como remitente '
         '(ej: farmacia@hospital-fuenlabrada.es).'],
    ]

    create_styled_table(doc, ['Campo', 'Descripción'], smtp_fields, [5, 12])

    doc.add_paragraph('')
    doc.add_paragraph(
        'Tras configurar los datos, puede pulsar "Enviar email de prueba" para verificar que '
        'la configuración es correcta. Se enviará un correo de prueba a la dirección del remitente.'
    )

    doc.add_heading('10.2. Email en Copia (CC)', level=2)

    doc.add_paragraph(
        'Opcionalmente, puede configurar una dirección de correo electrónico en copia (CC). '
        'Si se configura, todos los correos de notificación de préstamos se enviarán también a '
        'esta dirección. Esto es útil para que la secretaria del servicio o un responsable '
        'mantenga un registro de todos los préstamos notificados.'
    )

    doc.add_page_break()

    # ==================== 11. FLUJO DE TRABAJO ====================
    doc.add_heading('11. Flujo de Trabajo Completo', level=1)

    doc.add_heading('11.1. Flujo: Solicitamos Medicamento a Otro Hospital', level=2)

    doc.add_paragraph(
        'Este flujo se sigue cuando el HUF necesita un medicamento que no tiene disponible '
        'y lo solicita a otro hospital:'
    )

    flow_solicitar = [
        ('1. Identificación de la necesidad',
         'El farmacéutico identifica que se necesita un medicamento que no está disponible en el HUF '
         'y contacta telefónicamente con otro hospital para solicitar el préstamo.'),
        ('2. Registro en el sistema',
         'Una vez confirmada la disponibilidad por parte del hospital proveedor, el farmacéutico '
         'accede a MedLoan y crea un nuevo préstamo seleccionando:\n'
         '   — Tipo: "Solicitamos"\n'
         '   — Hospital: el hospital que nos prestará el medicamento\n'
         '   — Medicamento(s) y unidades\n'
         '   — Email del hospital proveedor\n'
         '   — Nombre del farmacéutico\n'
         '   — Notas relevantes (si las hay)'),
        ('3. Envío automático',
         'Al guardar, el sistema automáticamente:\n'
         '   — Genera el número de referencia (PREST-AAAA-NNNNN)\n'
         '   — Crea el PDF del préstamo con la imagen del hospital\n'
         '   — Envía el email al hospital proveedor con el PDF adjunto\n'
         '   — Si hay CC configurado, también envía copia'),
        ('4. Gestión en Farmatools',
         'El farmacéutico registra el préstamo en Farmatools (sistema de gestión de farmacia) '
         'y luego marca en MedLoan el toggle "Farmatools" como gestionado.'),
        ('5. Recepción del medicamento',
         'Cuando el medicamento llega físicamente al HUF, se verifica que coincide con lo registrado '
         'en el préstamo.'),
        ('6. Devolución',
         'Cuando el medicamento es devuelto al hospital proveedor, el farmacéutico marca el '
         'préstamo como "Devuelto" en MedLoan. El préstamo queda así completado.'),
        ('7. Seguimiento',
         'Si el préstamo no se devuelve en un tiempo razonable, aparecerá en la sección de '
         '"Préstamos atrasados" en Estadísticas y en el listado de "Pendientes de devolver".'),
    ]

    for title, desc in flow_solicitar:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}\n')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    doc.add_heading('11.2. Flujo: Nos Solicitan Medicamento desde Otro Hospital', level=2)

    doc.add_paragraph(
        'Este flujo se sigue cuando otro hospital contacta al HUF para solicitar un medicamento:'
    )

    flow_nos_solicitan = [
        ('1. Recepción de la solicitud',
         'El servicio de farmacia del HUF recibe una solicitud de préstamo de medicamento por parte '
         'de otro hospital (generalmente por teléfono o email).'),
        ('2. Verificación de disponibilidad',
         'El farmacéutico verifica la disponibilidad del medicamento solicitado y confirma al hospital '
         'solicitante.'),
        ('3. Registro en el sistema',
         'El farmacéutico accede a MedLoan y crea un nuevo préstamo seleccionando:\n'
         '   — Tipo: "Nos solicitan"\n'
         '   — Hospital: el hospital que nos solicita el medicamento\n'
         '   — Medicamento(s) y unidades\n'
         '   — Nombre del farmacéutico\n'
         '   — Notas relevantes\n\n'
         'Nota: Para este tipo de préstamo NO se envía email automático.'),
        ('4. Dispensación del medicamento',
         'Se prepara y entrega el medicamento al hospital solicitante.'),
        ('5. Gestión en Farmatools',
         'El farmacéutico registra la dispensación en Farmatools y marca el toggle correspondiente '
         'en MedLoan.'),
        ('6. Seguimiento de la devolución',
         'El préstamo permanece como "Pendiente de que nos devuelvan" hasta que el hospital solicitante '
         'devuelve el medicamento. Se puede generar un PDF del listado de pendientes para hacer '
         'seguimiento.'),
        ('7. Recepción de la devolución',
         'Cuando el hospital solicitante devuelve el medicamento, el farmacéutico marca el préstamo '
         'como "Devuelto" en MedLoan.'),
    ]

    for title, desc in flow_nos_solicitan:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}\n')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ==================== 12. ROLES Y RESPONSABILIDADES ====================
    doc.add_heading('12. Roles y Responsabilidades', level=1)

    doc.add_paragraph(
        'Aunque el sistema no gestiona usuarios individuales (todos acceden con la misma contraseña), '
        'sí existen distintos perfiles de uso dentro del Servicio de Farmacia:'
    )

    roles_data = [
        ['Farmacéutico\nresponsable de turno',
         '— Crear préstamos de ambos tipos.\n'
         '— Verificar que los datos son correctos antes de guardar.\n'
         '— Indicar su nombre en el campo "Farmacéutico".\n'
         '— Gestionar los estados (Farmatools, Devuelto).\n'
         '— Añadir notas relevantes al préstamo.\n'
         '— Descargar PDFs cuando sea necesario.',
         'Alta'],
        ['Secretaria del servicio\n/ Administrativo',
         '— Dar de alta nuevos hospitales con sus datos de contacto.\n'
         '— Mantener actualizado el catálogo de hospitales (emails, teléfonos, direcciones).\n'
         '— Generar listados de pendientes para seguimiento.\n'
         '— Exportar estadísticas para informes de gestión.\n'
         '— Realizar seguimiento de devoluciones pendientes.',
         'Media'],
        ['Jefe/a de Servicio\n/ Responsable',
         '— Consultar estadísticas y KPIs del servicio.\n'
         '— Revisar préstamos atrasados y tomar decisiones.\n'
         '— Exportar informes (PDF/Excel) para la dirección.\n'
         '— Supervisar la configuración del sistema.\n'
         '— Definir el umbral de días para considerar préstamos atrasados.',
         'Consulta y\nsupervisión'],
        ['Administrador\ntécnico',
         '— Configurar los parámetros SMTP de correo.\n'
         '— Realizar pruebas de envío de email.\n'
         '— Gestionar la contraseña de acceso al sistema.\n'
         '— Resolver incidencias técnicas.',
         'Configuración'],
    ]

    create_styled_table(doc, ['Perfil', 'Responsabilidades', 'Nivel de\ninteracción'],
                       roles_data, [4, 11, 2.5])

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 80, 0)
    p.add_run(
        'Todos los usuarios comparten el mismo nivel de acceso técnico al sistema. Las '
        'responsabilidades descritas son organizativas, no restricciones del sistema. Se recomienda '
        'que cada persona se ciña a las funciones asignadas a su perfil para mantener la coherencia '
        'y trazabilidad de los datos.'
    )

    doc.add_page_break()

    # ==================== 13. GLOSARIO ====================
    doc.add_heading('13. Glosario de Términos', level=1)

    glossary = [
        ['AEMPS', 'Agencia Española de Medicamentos y Productos Sanitarios. Organismo regulador de '
         'medicamentos en España.'],
        ['CC (Copia de Carbón)', 'Dirección de email que recibe una copia del correo electrónico enviado.'],
        ['CIMA', 'Centro de Información de Medicamentos Autorizados. Base de datos de la AEMPS que contiene '
         'información sobre todos los medicamentos autorizados en España.'],
        ['CN (Código Nacional)', 'Código numérico único que identifica cada presentación de un medicamento '
         'en el sistema sanitario español.'],
        ['Dashboard', 'Panel de control. Pantalla principal de la aplicación con resumen de KPIs y '
         'actividad reciente.'],
        ['Farmatools', 'Sistema de gestión de farmacia hospitalaria utilizado en el HUF para el control '
         'de stock, dispensación y gestión de medicamentos.'],
        ['HUF', 'Hospital Universitario de Fuenlabrada.'],
        ['KPI', 'Key Performance Indicator (Indicador Clave de Rendimiento). Métrica utilizada para '
         'evaluar el desempeño de un proceso.'],
        ['MedLoan', 'Nombre interno del sistema de gestión de préstamos de medicamentos.'],
        ['Nos solicitan', 'Tipo de préstamo en el que otro hospital solicita medicamentos al HUF. '
         'El HUF actúa como proveedor.'],
        ['Número de referencia', 'Identificador único del préstamo con formato PREST-AAAA-NNNNN '
         '(ej: PREST-2026-00001).'],
        ['PDF', 'Portable Document Format. Formato de documento utilizado para generar los informes '
         'y documentación de préstamos.'],
        ['Principio activo', 'Sustancia farmacológicamente activa de un medicamento.'],
        ['Presentación', 'Forma farmacéutica en la que se comercializa un medicamento (comprimidos, '
         'cápsulas, viales, etc.).'],
        ['SMTP', 'Simple Mail Transfer Protocol. Protocolo utilizado para el envío de correos electrónicos.'],
        ['Solicitamos', 'Tipo de préstamo en el que el HUF solicita medicamentos a otro hospital. '
         'El HUF actúa como receptor.'],
        ['Toggle', 'Interruptor visual en la interfaz que permite activar o desactivar un estado con un clic.'],
    ]

    create_styled_table(doc, ['Término', 'Definición'], glossary, [4.5, 12.5])

    doc.add_page_break()

    # ==================== 14. FAQ ====================
    doc.add_heading('14. Preguntas Frecuentes (FAQ)', level=1)

    faqs = [
        ('¿Qué hago si no encuentro un medicamento al crear un préstamo?',
         'Puede buscarlo en CIMA usando la opción "Buscar en CIMA" que aparece en el buscador de '
         'medicamentos. Si tampoco se encuentra en CIMA, puede crearlo manualmente pulsando '
         '"Crear nuevo" e introduciendo los datos disponibles.'),
        ('¿Se puede modificar un préstamo después de crearlo?',
         'Se pueden modificar las notas/observaciones y los estados (Farmatools y Devuelto). '
         'Los datos fundamentales del préstamo (hospital, medicamentos, unidades) no se pueden '
         'modificar una vez creado. Si ha cometido un error, deberá eliminar el préstamo y crear '
         'uno nuevo.'),
        ('¿Qué ocurre si el email no se envía correctamente?',
         'El préstamo se creará igualmente en el sistema. Se mostrará una notificación indicando '
         'que hubo un error en el envío del email. Verifique la configuración SMTP en la sección '
         'de Configuración y vuelva a intentarlo. Puede descargar el PDF manualmente y enviarlo '
         'por otros medios.'),
        ('¿Se pueden recuperar préstamos eliminados?',
         'No. La eliminación de préstamos es permanente e irreversible. Se recomienda precaución '
         'al usar la función de eliminar y siempre confirmar en el diálogo de confirmación.'),
        ('¿Cómo puedo ver solo los préstamos pendientes de devolución?',
         'En el listado de préstamos, utilice el filtro "Devolución" y seleccione "Pendiente". '
         'Puede combinarlo con el filtro de tipo para ver solo los pendientes de devolver o '
         'solo los pendientes de que nos devuelvan.'),
        ('¿Es necesario marcar "Farmatools" para marcar como "Devuelto"?',
         'No, ambos estados son independientes. Sin embargo, el flujo recomendado es primero '
         'gestionar en Farmatools y luego marcar como devuelto.'),
        ('¿Puedo usar el sistema desde el móvil?',
         'El sistema está diseñado principalmente para uso en escritorio, pero la interfaz es '
         'responsive y se adapta a dispositivos móviles. Se recomienda el uso en pantallas de al '
         'menos una tablet para una experiencia óptima.'),
        ('¿Cuánto tiempo se conservan los datos?',
         'Los datos se conservan indefinidamente en la base de datos. No existe un proceso automático '
         'de eliminación. Los préstamos, hospitales y medicamentos se mantienen para consulta histórica '
         'y estadísticas.'),
        ('¿Qué significa el formato del número de referencia?',
         'El formato PREST-AAAA-NNNNN se compone de:\n'
         '   — PREST: Prefijo fijo que indica "préstamo"\n'
         '   — AAAA: Año de creación (ej: 2026)\n'
         '   — NNNNN: Número secuencial que se incrementa automáticamente y se reinicia cada año'),
        ('¿Puedo exportar los datos de préstamos?',
         'Sí. Desde el listado de préstamos puede exportar en formato CSV. Desde la sección de '
         'Estadísticas puede exportar informes completos en PDF o Excel (XLSX).'),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f'P: {q}')
        run.bold = True
        run.font.color.rgb = RGBColor(13, 148, 136)
        p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        run = p.add_run('R: ')
        run.bold = True
        p.add_run(a)
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ==================== 15. SOPORTE ====================
    doc.add_heading('15. Soporte y Contacto', level=1)

    doc.add_paragraph(
        'Para cualquier incidencia, duda o sugerencia sobre el sistema MedLoan, póngase en contacto '
        'con el equipo técnico responsable del mantenimiento del sistema.'
    )

    doc.add_heading('Tipos de incidencias:', level=3)

    support_types = [
        ('Incidencia técnica',
         'El sistema no funciona, muestra errores, no se pueden crear préstamos, etc. '
         'Contacte con el administrador técnico.'),
        ('Problema con el envío de email',
         'Los correos no se envían o llegan con errores. Verifique primero la configuración SMTP '
         'y pruebe con "Enviar email de prueba". Si persiste, contacte con el administrador.'),
        ('Solicitud de cambio de contraseña',
         'Contacte con el administrador técnico para solicitar un cambio de la contraseña de acceso.'),
        ('Datos incorrectos en hospitales o medicamentos',
         'Un farmacéutico con acceso puede corregir los datos directamente desde las secciones '
         'correspondientes.'),
        ('Sugerencias de mejora',
         'Cualquier sugerencia para mejorar el sistema puede comunicarse al equipo técnico.'),
    ]

    for title, desc in support_types:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    # ==================== PIE FINAL ====================
    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(13, 148, 136)

    add_formatted_paragraph(doc,
        'Documento elaborado para el Servicio de Farmacia del Hospital Universitario de Fuenlabrada.',
        font_size=9, italic=True, color=(130, 130, 130),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_formatted_paragraph(doc,
        'Versión 1.0 — Febrero 2026',
        font_size=9, italic=True, color=(130, 130, 130),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_formatted_paragraph(doc,
        'Este documento es de uso interno y confidencial.',
        font_size=9, italic=True, color=(130, 130, 130),
        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ==================== GUARDAR ====================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Protocolo_Uso_MedLoan_HUF.docx'
    )
    doc.save(output_path)
    print(f'Informe generado exitosamente: {output_path}')
    return output_path

if __name__ == '__main__':
    generate_report()
